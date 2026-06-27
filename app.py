import os
import re
import io
from datetime import datetime
from flask import Flask, render_template_string, request, send_file

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

app = Flask(__name__)

# --- WORD DOCUMENT GENERATION (unchanged) ---
def create_element(name):
    return OxmlElement(name)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = create_element('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = create_element(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_table_styles(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="A6B1E1"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="A6B1E1"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(19, 60, 85)
    return p

def format_cell_text(cell, bold=False, text="", color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="133C55"/></w:pBdr>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F8"/>')
    pPr.append(pBdr)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(40, 40, 40)

def extract_table_names(sql_text):
    pattern = r'(?:from|update|into)\s+([a-zA-Z0-9_]+)'
    matches = re.findall(pattern, sql_text, re.IGNORECASE)
    forbidden = {'set', 'where', 'join', 'inner', 'left', 'select', 'and'}
    unique_tables = sorted(list(set([t.lower() for t in matches if t.lower() not in forbidden])))
    if not unique_tables:
        return "State Database Tables"
    return "\n".join([f"{i+1}) {table}" for i, table in enumerate(unique_tables)])

def generate_docx(data):
    doc = docx.Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    r_title = p_title.add_run("Issue Analysis & Resolution Report\nby Bidar Dist. Bhoomi Consultant")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(11, 37, 69)
    
    # 1) CLS Details
    add_styled_heading(doc, "1) CLS Details:")
    t1 = doc.add_table(rows=2, cols=4)
    apply_table_styles(t1)
    headers1 = ["Date", "District", "Consultant Name", "Call Log System Number"]
    for i, h in enumerate(headers1):
        set_cell_shading(t1.cell(0, i), "0B2545")
        set_cell_margins(t1.cell(0, i))
        format_cell_text(t1.cell(0, i), bold=True, text=h, color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    data1 = [data['report_date'], "Bidar", "Aravind Badagi", data['cls_number']]
    for i, d in enumerate(data1):
        set_cell_margins(t1.cell(1, i))
        format_cell_text(t1.cell(1, i), text=d, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    t1.columns[0].width = Inches(2.0)
    t1.columns[1].width = Inches(2.0)
    t1.columns[2].width = Inches(2.5)
    t1.columns[3].width = Inches(2.89)
    
    # 2) Issue Details
    add_styled_heading(doc, "2) Details of the Issue:")
    t2 = doc.add_table(rows=2, cols=4)
    apply_table_styles(t2)
    headers2 = ["Project Name", "Module Name", "Issue Type\n(Administrative/Technical)", "Issue Description"]
    for i, h in enumerate(headers2):
        set_cell_shading(t2.cell(0, i), "0B2545")
        set_cell_margins(t2.cell(0, i))
        format_cell_text(t2.cell(0, i), bold=True, text=h, color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
    format_cell_text(t2.cell(1, 0), text="Bhoomi", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell_text(t2.cell(1, 1), text="Bhoomi", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell_text(t2.cell(1, 2), text=data.get('issue_type', 'Technical'), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_margins(t2.cell(1, 3))
    issue_lines = []
    if data.get('issue_desc'):
        issue_lines.append(data['issue_desc'])
    meta_fields = [('App No', data.get('app_no')), ('Tran No', data.get('tran_no')),
                   ('Year', data.get('year')), ('Land Code', data.get('land_code')),
                   ('Season', data.get('season'))]
    for label, value in meta_fields:
        if value and value.strip():
            issue_lines.append(f"{label}: {value}")
    issue_text = "\n".join(issue_lines) if issue_lines else "No description provided"
    format_cell_text(t2.cell(1, 3), text=issue_text)
    for i in range(4): 
        set_cell_margins(t2.cell(1, i))
    t2.columns[0].width = Inches(1.5)
    t2.columns[1].width = Inches(1.5)
    t2.columns[2].width = Inches(1.8)
    t2.columns[3].width = Inches(4.59)
    
    tables_list = extract_table_names(data['select_block'] + "\n" + data['update_block'])
    
    # 3) Findings and Analysis
    add_styled_heading(doc, "3) Findings and Analysis:")
    t3 = doc.add_table(rows=2, cols=4)
    apply_table_styles(t3)
    headers3 = ["Data Tables involved", "Findings and Analysis", "Root Cause for the issue", "Solution Provided"]
    for i, h in enumerate(headers3):
        set_cell_shading(t3.cell(0, i), "0B2545")
        set_cell_margins(t3.cell(0, i))
        format_cell_text(t3.cell(0, i), bold=True, text=h, color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
    set_cell_margins(t3.cell(1, 0))
    format_cell_text(t3.cell(1, 0), text=tables_list)
    set_cell_margins(t3.cell(1, 1))
    format_cell_text(t3.cell(1, 1), text=data['findings_desc'])
    set_cell_margins(t3.cell(1, 2))
    format_cell_text(t3.cell(1, 2), text=data['root_cause'])
    set_cell_margins(t3.cell(1, 3))
    format_cell_text(t3.cell(1, 3), text=data['sol_provided'])
    t3.columns[0].width = Inches(2.0)
    t3.columns[1].width = Inches(2.5)
    t3.columns[2].width = Inches(2.0)
    t3.columns[3].width = Inches(2.89)
    
    add_styled_heading(doc, "4) \"Select\" Query for all the related/analyzed tables:")
    add_code_block(doc, data['select_block'])
    
    add_styled_heading(doc, "5) \"Update\" Query for the solution:")
    add_code_block(doc, data['update_block'])
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ===================================================================
#  COMPLETE HTML TEMPLATE WITH ALL VILLAGES EMBEDDED
# ===================================================================
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Bidar District · Revenue Command Center</title>
    <style>
        /* ----- Reset & Base ----- */
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Segoe UI', Roboto, system-ui, sans-serif; background: #f0f4f8; color: #1e293b; padding: 20px; min-height: 100vh; }
        .container { max-width: 1440px; margin: 0 auto; }
        .header { background: linear-gradient(135deg, #0b1a33 0%, #1a3a5c 60%, #2a5a7a 100%); border-radius: 24px; padding: 24px 32px; margin-bottom: 30px; box-shadow: 0 12px 40px rgba(10,30,60,0.25); border-bottom: 5px solid #d4a843; position: relative; overflow: hidden; }
        .header::before { content: ''; position: absolute; top: -50%; right: -10%; width: 300px; height: 300px; background: radial-gradient(circle, rgba(212,168,67,0.08) 0%, transparent 70%); border-radius: 50%; }
        .header-grid { display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 16px; position: relative; z-index: 1; }
        .header-left { display: flex; align-items: center; gap: 18px; }
        .emblem { width: 70px; height: 70px; background: #f8f4ea; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 34px; font-weight: 700; color: #0b1a33; box-shadow: 0 4px 12px rgba(0,0,0,0.2); border: 3px solid #d4a843; }
        .header-titles .gov-tag { font-size: 12px; font-weight: 700; letter-spacing: 2.5px; color: #d4a843; text-transform: uppercase; }
        .header-titles .main-title { font-size: 28px; font-weight: 700; color: #ffffff; text-shadow: 0 2px 8px rgba(0,0,0,0.15); line-height: 1.2; }
        .header-titles .sub-title { font-size: 18px; font-weight: 400; color: #cbd5e1; letter-spacing: 0.3px; }
        .header-right { text-align: right; background: rgba(255,255,255,0.06); padding: 8px 24px; border-radius: 40px; backdrop-filter: blur(4px); border: 1px solid rgba(212,168,67,0.2); }
        .header-right .dept { font-size: 14px; font-weight: 600; color: #d4a843; letter-spacing: 0.5px; }
        .header-right .location { font-size: 20px; font-weight: 700; color: #ffffff; letter-spacing: 0.5px; }
        .header-right .location small { font-weight: 400; font-size: 14px; color: #94a3b8; }
        .card { background: white; border-radius: 20px; padding: 24px; border: 1px solid #e2e8f0; box-shadow: 0 6px 20px rgba(0,0,0,0.03); transition: box-shadow 0.2s; }
        .card:hover { box-shadow: 0 8px 30px rgba(0,0,0,0.06); }
        .card-title { display: flex; align-items: center; gap: 12px; font-size: 16px; font-weight: 700; color: #0b1a33; border-bottom: 2px solid #e2e8f0; padding-bottom: 12px; margin-bottom: 20px; text-transform: uppercase; letter-spacing: 0.5px; }
        .card-title .step-badge { background: #0b1a33; color: #d4a843; padding: 2px 14px; border-radius: 30px; font-size: 12px; font-weight: 700; }
        .dashboard-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 24px; margin-bottom: 30px; }
        @media (max-width: 1100px) { .dashboard-grid { grid-template-columns: 1fr; } }
        .filter-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(160px, 1fr)); gap: 14px; margin-bottom: 16px; }
        .filter-group label { display: block; font-size: 11px; font-weight: 700; color: #334155; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 4px; }
        .filter-group select, .filter-group input { width: 100%; padding: 8px 12px; font-size: 14px; border: 1.5px solid #cbd5e1; border-radius: 10px; background: #f8fafc; font-family: inherit; transition: 0.2s; outline: none; color: #1e293b; }
        .filter-group select:focus, .filter-group input:focus { border-color: #0b1a33; box-shadow: 0 0 0 3px rgba(11,26,51,0.1); }
        .table-wrapper { max-height: 260px; overflow-y: auto; border: 1.5px solid #e2e8f0; border-radius: 12px; margin-top: 12px; }
        .table-wrapper table { width: 100%; border-collapse: collapse; font-size: 13px; }
        .table-wrapper th { background: #f1f5f9; padding: 10px 12px; text-align: left; position: sticky; top: 0; z-index: 2; color: #0b1a33; font-weight: 700; border-bottom: 2px solid #cbd5e1; }
        .table-wrapper td { padding: 8px 12px; border-bottom: 1px solid #e2e8f0; cursor: pointer; }
        .table-wrapper tr:hover td { background: #f8fafc; }
        .village-link { color: #0b1a33; font-weight: 600; text-decoration: none; border-bottom: 2px solid #d4a843; transition: 0.15s; cursor: pointer; }
        .village-link:hover { color: #d4a843; border-bottom-color: #0b1a33; }
        .no-results { text-align: center; padding: 30px; color: #64748b; font-style: italic; }
        .checkbox-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(150px, 1fr)); gap: 6px; max-height: 200px; overflow-y: auto; padding: 12px; border: 1.5px solid #e2e8f0; border-radius: 12px; background: #f8fafc; margin-bottom: 16px; }
        .checkbox-grid label { display: flex; align-items: center; gap: 6px; font-size: 13px; color: #1e293b; cursor: pointer; padding: 2px 4px; border-radius: 4px; transition: 0.1s; }
        .checkbox-grid label:hover { background: #e2e8f0; }
        .checkbox-grid .all-tables { grid-column: 1 / -1; font-weight: 700; padding: 6px 0 10px 0; border-bottom: 2px solid #cbd5e1; margin-bottom: 6px; color: #0b1a33; }
        .checkbox-grid input[type="checkbox"] { width: 16px; height: 16px; accent-color: #0b1a33; cursor: pointer; flex-shrink: 0; }
        .btn-group { display: flex; gap: 10px; flex-wrap: wrap; align-items: center; }
        .btn { padding: 8px 20px; border-radius: 10px; font-weight: 600; font-size: 13px; cursor: pointer; border: none; transition: all 0.2s; font-family: inherit; display: inline-flex; align-items: center; gap: 6px; }
        .btn-primary { background: #0b1a33; color: white; box-shadow: 0 4px 12px rgba(11,26,51,0.2); }
        .btn-primary:hover { transform: translateY(-1px); box-shadow: 0 6px 20px rgba(11,26,51,0.3); }
        .btn-secondary { background: #e2e8f0; color: #1e293b; border: 1px solid #cbd5e1; }
        .btn-secondary:hover { background: #cbd5e1; }
        .btn-success { background: #0b1a33; color: #d4a843; box-shadow: 0 4px 12px rgba(11,26,51,0.2); width: 100%; justify-content: center; padding: 12px; font-size: 15px; font-weight: 700; }
        .btn-success:hover { background: #1a3a5c; color: #ffd966; }
        .btn-sm { padding: 5px 14px; font-size: 12px; border-radius: 8px; }
        .query-output { background: #0b1a33; border-radius: 12px; padding: 16px; color: #e2e8f0; font-family: 'Consolas', 'Courier New', monospace; font-size: 12px; max-height: 200px; overflow-y: auto; white-space: pre-wrap; word-break: break-all; border-left: 4px solid #d4a843; margin-top: 12px; line-height: 1.8; }
        .query-output .empty { color: #94a3b8; font-style: italic; }
        .query-output .highlight { color: #ffd966; }
        .query-output .table-name { color: #7dd3fc; font-weight: 600; }
        .query-output .keyword { color: #f472b6; }
        .query-output .number { color: #a78bfa; }
        .query-output .string { color: #fbbf24; }
        .code-stats { font-size: 13px; font-weight: 600; color: #0b1a33; background: #f1f5f9; padding: 4px 16px; border-radius: 20px; border: 1px solid #cbd5e1; white-space: nowrap; }
        .form-group { margin-bottom: 16px; }
        .form-group label { display: block; font-size: 12px; font-weight: 700; color: #334155; text-transform: uppercase; letter-spacing: 0.4px; margin-bottom: 4px; }
        .form-group textarea, .form-group input, .form-group select { width: 100%; padding: 8px 12px; font-size: 14px; border: 1.5px solid #cbd5e1; border-radius: 10px; background: #f8fafc; font-family: inherit; transition: 0.2s; outline: none; color: #1e293b; }
        .form-group textarea:focus, .form-group input:focus, .form-group select:focus { border-color: #0b1a33; box-shadow: 0 0 0 3px rgba(11,26,51,0.1); }
        .form-group textarea { min-height: 60px; resize: vertical; }
        .form-group .code-input { font-family: 'Consolas', 'Courier New', monospace; background: #0b1a33; color: #e2e8f0; border-color: #1e3a5a; min-height: 80px; }
        .form-row { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }
        @media (max-width: 600px) { .form-row { grid-template-columns: 1fr; } }
        .radio-group { display: flex; flex-wrap: wrap; gap: 6px 18px; background: #f1f5f9; padding: 8px 16px; border-radius: 30px; border: 1px solid #cbd5e1; margin-bottom: 12px; }
        .radio-group label { display: flex; align-items: center; gap: 6px; font-size: 13px; font-weight: 600; color: #1e293b; cursor: pointer; }
        .radio-group input[type="radio"] { accent-color: #0b1a33; width: 15px; height: 15px; cursor: pointer; }
        .footer { margin-top: 30px; text-align: center; padding: 16px; color: #64748b; font-size: 14px; border-top: 1px solid #e2e8f0; }
        .footer a { color: #0b1a33; text-decoration: none; font-weight: 600; }
        .footer a:hover { text-decoration: underline; color: #d4a843; }
        @media (max-width: 768px) {
            .header-grid { flex-direction: column; align-items: flex-start; }
            .header-right { text-align: left; width: 100%; }
            .header-titles .main-title { font-size: 22px; }
            .header-titles .sub-title { font-size: 16px; }
            .emblem { width: 56px; height: 56px; font-size: 26px; }
            .card { padding: 16px; }
            .filter-grid { grid-template-columns: 1fr 1fr; }
            .checkbox-grid { grid-template-columns: 1fr 1fr; max-height: 150px; }
        }
        @media (max-width: 480px) {
            .filter-grid { grid-template-columns: 1fr; }
            .checkbox-grid { grid-template-columns: 1fr; }
            .btn-group { flex-direction: column; width: 100%; }
            .btn-group .btn { width: 100%; justify-content: center; }
        }
    </style>
</head>
<body>

<div class="container">

    <!-- ===== HEADER ===== -->
    <header class="header">
        <div class="header-grid">
            <div class="header-left">
                <span class="emblem" aria-hidden="true">🏛️</span>
                <div class="header-titles">
                    <h1 class="main-title">Deputy Commissioner Office</h1>
                    <h2 class="sub-title">Revenue Department, Bidar</h2>
                </div>
            </div>
            <div class="header-right">
                <div class="dept"><span aria-hidden="true">📋</span> District Administration</div>
                <div class="location">Bidar <small>· Village Operations Panel</small></div>
            </div>
        </div>
    </header>

    <!-- ===== DASHBOARD ===== -->
    <div class="dashboard-grid">

        <!-- LEFT COLUMN: Village Explorer + Query Generator -->
        <div>

            <!-- STEP 1: Village Explorer -->
            <div class="card">
                <div class="card-title">
                    <span class="step-badge">1</span>
                    Village Explorer &amp; Code Lookup
                </div>

                <div class="filter-grid">
                    <div class="filter-group">
                        <label>District</label>
                        <select id="districtSelect"><option value="">All</option></select>
                    </div>
                    <div class="filter-group">
                        <label>Taluk</label>
                        <select id="talukSelect"><option value="">All</option></select>
                    </div>
                    <div class="filter-group">
                        <label>Hobli</label>
                        <select id="hobliSelect"><option value="">All</option></select>
                    </div>
                    <div class="filter-group">
                        <label>Village</label>
                        <select id="villageSelect"><option value="">All</option></select>
                    </div>
                </div>

                <div class="filter-group">
                    <label>🔍 Quick Search Village Name</label>
                    <input type="text" id="villageSearch" placeholder="Type to filter villages...">
                </div>

                <div class="table-wrapper">
                    <table>
                        <thead>
                            <tr><th>District</th><th>Taluk</th><th>Hobli</th><th>Village Name</th></tr>
                        </thead>
                        <tbody id="tableBody"><tr><td colspan="4" class="no-results">Loading directory…</td></tr></tbody>
                    </table>
                </div>
            </div>

            <!-- STEP 2: Query Generator -->
            <div class="card">
                <div class="card-title">
                    <span class="step-badge">2</span>
                    System Query Generator
                </div>

                <div style="display:flex; flex-wrap:wrap; gap:10px; margin-bottom:12px;">
                    <div class="filter-group" style="flex:1 1 160px;">
                        <label>📄 Appl No</label>
                        <input type="text" id="applNoInput" placeholder="e.g. 123/2024">
                    </div>
                    <div class="filter-group" style="flex:1 1 160px;">
                        <label>🔢 Tran No</label>
                        <input type="text" id="tranNoInput" placeholder="e.g. 45/2024">
                    </div>
                    <div class="filter-group" style="flex:1 1 160px;">
                        <label>🗺️ Land Code</label>
                        <input type="text" id="landCodeInput" placeholder="e.g. LAND123">
                    </div>
                    <div class="filter-group" style="flex:1 1 160px;">
                        <label>📅 Year Code</label>
                        <input type="text" id="yearCodeInput" placeholder="e.g. 2024">
                    </div>
                </div>

                <label style="font-size:12px;font-weight:700;color:#334155;text-transform:uppercase;letter-spacing:0.4px;display:block;margin-bottom:6px;">
                    📋 Select Tables
                </label>
                <div class="radio-group" id="tableGroupRadio">
                    <label><input type="radio" name="tableGroup" value="all" checked> All</label>
                    <label><input type="radio" name="tableGroup" value="application"> Application</label>
                    <label><input type="radio" name="tableGroup" value="transaction"> Transaction</label>
                    <label><input type="radio" name="tableGroup" value="history"> History</label>
                    <label><input type="radio" name="tableGroup" value="rtc"> RTC</label>
                    <label><input type="radio" name="tableGroup" value="gsc"> GSC</label>
                </div>
                <div id="tableSelectContainer" class="checkbox-grid"></div>

                <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:10px;">
                    <div class="btn-group">
                        <button class="btn btn-secondary btn-sm" id="clearAllBtn">✖ Reset</button>
                        <button class="btn btn-primary btn-sm" id="submitBtn">⚡ Compile Queries</button>
                    </div>
                    <span class="code-stats" id="codeStats">📌 Codes: —</span>
                </div>

                <div style="margin-top:14px;">
                    <label style="font-size:12px;font-weight:700;color:#334155;text-transform:uppercase;letter-spacing:0.4px;display:block;margin-bottom:6px;">
                        📄 Generated Queries
                    </label>
                    <div class="query-output" id="queryGrid">
                        <span class="empty">Select filters and click Compile</span>
                    </div>
                </div>
            </div>
        </div>

        <!-- RIGHT COLUMN: Issue Report Generator -->
        <div>
            <div class="card" style="height:100%;">
                <div class="card-title">
                    <span class="step-badge">3</span>
                    Issue Report Generator
                </div>

                <form method="POST" action="/generate" id="reportForm">
                    <div class="form-row">
                        <div class="form-group">
                            <label>📅 Report Date</label>
                            <input type="text" name="report_date" value="{{ current_date }}" required>
                        </div>
                        <div class="form-group">
                            <label>📞 CLS Number</label>
                            <input type="text" name="cls_number" placeholder="Optional reference">
                        </div>
                    </div>

                    <div class="form-row">
                        <div class="form-group">
                            <label>📄 App No</label>
                            <input type="text" name="app_no" id="reportAppNo" placeholder="Optional">
                        </div>
                        <div class="form-group">
                            <label>🧾 Tran No</label>
                            <input type="text" name="tran_no" id="reportTranNo" placeholder="Optional">
                        </div>
                    </div>

                    <div class="form-row">
                        <div class="form-group">
                            <label>📆 Year</label>
                            <input type="text" name="year" id="reportYear" placeholder="Optional">
                        </div>
                        <div class="form-group">
                            <label>🏞️ Land Code</label>
                            <input type="text" name="land_code" id="reportLandCode" placeholder="Optional">
                        </div>
                    </div>

                    <div class="form-row">
                        <div class="form-group">
                            <label>🌦️ Season</label>
                            <input type="text" name="season" id="reportSeason" placeholder="Optional">
                        </div>
                        <div class="form-group">
                            <label>📋 Issue Type</label>
                            <select name="issue_type" required>
                                <option value="Technical">Technical</option>
                                <option value="Administrative">Administrative</option>
                            </select>
                        </div>
                    </div>

                    <div class="form-group">
                        <label>📝 Issue Description</label>
                        <textarea id="form_issue_desc" name="issue_desc" placeholder="Describe the issue in detail..." required></textarea>
                    </div>

                    <div class="form-group">
                        <label>🔍 Findings &amp; Analysis</label>
                        <textarea name="findings_desc" placeholder="Document your findings and analysis..." required></textarea>
                    </div>

                    <div class="form-group">
                        <label>🎯 Root Cause</label>
                        <input type="text" name="root_cause" placeholder="e.g., Data mismatch between tables" required>
                    </div>

                    <div class="form-group">
                        <label>✅ Solution Provided</label>
                        <textarea name="sol_provided" placeholder="Describe the solution implemented..." required></textarea>
                    </div>

                    <div class="form-row">
                        <div class="form-group">
                            <label>🔎 SELECT Query</label>
                            <textarea id="form_select_block" name="select_block" class="code-input" placeholder="SELECT * FROM table WHERE..." required></textarea>
                        </div>
                        <div class="form-group">
                            <label>✏️ UPDATE Query</label>
                            <textarea id="form_update_block" name="update_block" class="code-input" placeholder="UPDATE table SET..." required></textarea>
                        </div>
                    </div>

                    <button type="submit" class="btn btn-success">💾 Generate &amp; Download Word Report</button>
                </form>
            </div>
        </div>
    </div>

    <!-- ===== FOOTER ===== -->
    <div class="footer">
        Copyright © 2025 Bhoomi Monitoring Cell · Government of Karnataka · Bidar District · Revenue Department
        <br>
        Developed &amp; Designed By:
        <a href="https://arvindbadagi.github.io/analytics/" target="_blank">Arvind Badagi</a>
    </div>
</div>

<script>
    // ================================================================
    //  FULL DATASET – All villages (complete)
    // ================================================================
    const rawRowsOriginal = [
        // BASAVAKALYAN TALUK
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","BASAVAKALYANA (1)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","YADALAPURA (2)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","SHIVAPURA (3)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","TRIPURANKA (4)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","TALABOGA (5)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","PRATAPURA (6)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","MORAKHANDI (7)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","GOWRA (8)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","JANAPURA (9)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","GOKULA (10)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","KITTA (11)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","DHANNURA K (12)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","KHANDALA (13)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","KHANAPURA K (14)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","NEELAKANTA (15)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","LAHESHVARA (16)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","NARAYANAPURA (17)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","BASAVAKALYANA (1)","BETABALAKUNDA (18)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","MANTALA (1)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","ILYALA (2)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","GUNDURA (3)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","ALAGUDA (4)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","VATTARAGA (5)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","HIPPARAGA GHATA (6)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","HONNALLI (7)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","CHITTA KALLADEVA (8)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","URKI (9)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","MANNALLI (10)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","KAMALELUDI (11)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","ATALAPURA (12)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","SASTAPURA (13)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","HANDRIYALA RU (14)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","JAJANAMUGALI (15)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","GHOTALA (16)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","RAMATIRTHAKH (17)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","UMAPURA (18)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","HALLI (19)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","CHANDAKAPURA (20)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MANTALA (2)","MIRJAPURA (21)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","RAJESHVARA (1)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","RAJOLA (2)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","ISLAMPURA (3)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","HANDRALAKE (4)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","LINGADALLI (5)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","NEERAGUDI (6)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","KOWDIYALA  S (7)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","ROLA (8)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","KOWDIYALARU (9)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","KHERDABH (10)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","SADALAPURA (11)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","DHANNURARU (12)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","PANDARAGERA (13)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","YARANDI (14)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","MANGALURA (15)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","YARABAGA (16)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","TADOLA (17)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","RAJESHVARA (3)","GHOGGA (18)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","MUDABI (1)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","KHANAPURA BH (2)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","BAGADHURI (3)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","HATYALA (4)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","HARAKUDA (5)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","SARAJAVALAGA (6)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","JANAVADA (7)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","HIPPARAGA BH (8)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","MAISALAGA (9)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","KALAKHORA (10)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","HIRANAGANVA (11)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","SUNGATANA (12)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","EKALURA (13)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","GADALEGANVA BH (14)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","YALAVANTI (15)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","KHERDA KH (16)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","CHIKANAGANVA (17)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","SIRAGAPURA (18)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","KINNI (19)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","YADALAGUNDI (20)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","SAIDAPURA (21)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","MUDABI (4)","HAMUNAGARA (22)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","KOHINURA (1)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","LADAVANTI (2)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","EKKAMBA (3)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","DHAMURI (4)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","UJALAMBA (5)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","CHITTAKOTA (6)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","BHAKANALA (7)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","BETAGERA (8)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","GADALEGANVA (9)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","SIRAGURA (10)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","RAMATIRTHAD (11)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","GILIGILI (12)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","HATTARAGA S (13)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","CHITTAKOTAK (14)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","BHOSAGA (15)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","MANAKHEDA (16)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","ATTURA (17)"],
        ["BIDAR (5)","BASAVAKALYAN (1)","KOHINURA (5)","SIRURI (18)"],
        // BHALKI TALUK
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","BHALKI (1)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","BAJULAGA (2)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KORURA (3)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KARADIYALA (4)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KADALABADA (5)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","GUBYALA (6)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","DHARAJAVADI (8)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","TALAVADA KH (9)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","SAIDAPURAVADI (10)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","JOLADABAKA (11)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","SIDDESHVARA (12)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","MARURA (13)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","CHIKALACHANDA (14)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KUNTESHIRSHI (15)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","MADAKATTI (16)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KALAVADI (17)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","HARANALA (18)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","DADAGI (19)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","RAMATIRTHAVADI (20)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","EKALASAPURAVADI (21)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","GORACHINCHOLI (22)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","NIDEBAN (23)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","ANANDAVADI (24)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KOTAGERA (25)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","KOTAGYALAVADI (26)"],
        ["BIDAR (5)","BHALKI (2)","BHALKI (1)","GANESHAPURAVADI (27)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","SAYAGANVA (1)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","KHUDAVANDAPURA (2)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","INCHURA (3)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","KESARA JAVALAGA (4)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","SHREEMALI (5)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","JAMAKHANDI (6)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","VANJARAKHEDA (7)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","KONGALI (8)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","RU GOWNDAGANVA (9)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","KUNTEGANVA (10)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","LADA (11)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","HALASI L (12)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","JEERAGYALA (13)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","VAGALAGANVA (14)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","GUNJARAGA (15)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","ATTARAGA (16)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","ALAVAYI (17)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","NARADASANGAMA (18)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","MEHAKARA (19)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","MANIKESHVARA (20)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","TUGANVA H (21)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","BOLEGANVA (22)"],
        ["BIDAR (5)","BHALKI (2)","SAYAGANVA (2)","ELAMAVADI (23)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","NITTURA BH (1)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","AMBESANGAVI (2)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","HAJANALA (3)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","GORANALA (4)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","BEERIKH (5)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","NAGARALA (6)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","MURALA (7)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","KUDLI (8)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","HUPALA (9)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","KERURA (10)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","AHMADABADA (11)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","JAINAPURA (12)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","BALURA (13)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","KOTAGYALA (14)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","CHANDAPURA (15)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","KHASHAMAPURA (16)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","DONGARAGI (17)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","VALASANGA (18)"],
        ["BIDAR (5)","BHALKI (2)","NITTURABH (3)","SIDDAPURA VADI (19)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","KHATAKA CHINCHOLI (1)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","UCHYA (2)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","BYALAHALLI W (3)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","VARAVATTI (4)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","DAVARAGANVA (6)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","TUGANVACH (7)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","NELVADA (8)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","MORAMBI (9)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","HONNALLI (10)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","HALAGORTA (11)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","CHALAKAPURA (12)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","NAVADAGI (13)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","KAPALAPURA (14)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","MAVINAHALLI (15)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","ENAKURA (16)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","BARADAPURA (17)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","KURABA KHELAGI (18)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","MASHIMADU (19)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","CHITTA (20)"],
        ["BIDAR (5)","BHALKI (2)","KHATAKA CHINCHULI (4)","SIKINDRABADA VADI (21)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","LAKHANAGANVA (1)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","ALANDI (2)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","KALASADALA (3)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","HACHCHIKAMATA (4)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","DONAGAPURA (5)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","KAKANALA (6)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","METIMELAKUNDA (7)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","BHATAMBRA (8)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","BEERIBH (9)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","HUNAJIA (10)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","LANJAVADA (12)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","BHATASANGAVI (13)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","SHIVANI (14)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","TALAVADAM (15)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","TELAGANVA (16)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","SOMAPURA (17)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","ATANURA (18)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","KALASARATUGANVA (19)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","JAYAGANVA (20)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","TAMAGYALA (21)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","SHAMASHERAPURA (22)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","METIMELAKUNDA VADI (25)"],
        ["BIDAR (5)","BHALKI (2)","LAKHANAGANVA (5)","KASARATUGANVA VADI (26)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","HALABARGA (1)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","KANAJI (2)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","TEGAMPURA (3)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","HALAHIPPARAGA (4)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","KOSAMA (5)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","NAGURAKE (6)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","ALIYABADA (7)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","TARANALLI (8)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","HUNAJEE K (9)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","GODIHIPPARGA (10)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","DHANNURAH (11)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","KONAMELAKUNDA (12)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","KHANAPURA (13)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","RUDANURA (14)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","MALACHAPURA (15)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","JANTI (16)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","NELAGI (17)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","BYALAHALLIKH (18)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","KAMALAPURA (19)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","NIRAMANAHALLI (20)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","SEVANAGARA (21)"],
        ["BIDAR (5)","BHALKI (2)","HALABARGA (6)","HALAHALLI K (22)"],
        // AURAD TALUK
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","OWRADABI (1)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","MAMADAPURA (2)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","NARAYANAPURA (3)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","GANESHAPURE (4)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","ALLAPURA (5)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","NARASINHAPURA (6)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","EKALARA (7)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","TULAJAPURA (8)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","BORALA (9)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","BADALAGANVA (10)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","DUDUKANALA (11)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","VANAMARAPALLI (12)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","MUNGANALA (13)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","KAPPIKERI (14)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","HASSIKERI (15)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","KHANDIKERI (16)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","HULLYALA (17)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","EKAMBA (18)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","JAMALAPURA (19)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","KOLLURA (20)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","JONNIKERI (21)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","MAHADONAGANVA (22)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","TEGAMPURA (23)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","GOWNDAGANVA (24)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","BONTI (26)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","HANGARAGABI (27)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","HOKRANA (28)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","KHERDABI (29)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","KARAKYALA (30)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","DONGARAGAVA (31)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","BANDARA KUMATA (32)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","LINGI (33)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","BAVALAGANVA (34)"],
        ["BIDAR (5)","AURAD (3)","OWRADABI (1)","SAVARAGANVA (35)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","CHINTAKI (1)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","BELDALA (2)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","JOJANA (3)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","NAGURAN (4)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","LINGADALSHIKH (5)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","KARANJIBI (6)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","KARANJIKE (7)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","MANURAKE (8)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","RAYAPALLI (9)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","NAGANAPALLI (10)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","LINGADALLI J (11)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","SUNDALA (12)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","KHASEMPURAA (13)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","UJANI (14)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","SUNKANALA (15)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","GUDAPALLI (16)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","NAGAMARAPALLI (17)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","NANDYALA (18)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","JAKANALA (19)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","ITAGYALA (20)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","ENAGUNDA (21)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","MEDAPALLI (22)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","NAGUREM (23)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","CHIKLIJE (24)"],
        ["BIDAR (5)","AURAD (3)","CHINTAKI (2)","BARADAPURA (25)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","SANTAPURA (1)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","JIRGABI (2)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","JEERGAKE (3)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","JAMBAGI (4)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","MAHARAJAVADI (5)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","SHEMBELLI (6)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","CHATANALA (7)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","GADIKUSANURA (8)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","PASHAPURA (9)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","BALURAJE (10)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","KOWTABI (11)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","KOWTAKE (12)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","BELUREN (13)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","VALLEPURA (14)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","ALURAKE (15)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","SORALLI (16)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","KOWDAGANVA (17)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","ALURABI (18)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","MUSTAPURA (19)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","BORGIJE (20)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","KANDAGULA (22)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","KHANAPURA (23)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","VADAGANVAD (24)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","KHASEMPURA B (25)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","RAKSHSHALA K (26)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","MASKALA (27)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","NAGURA B (28)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","BELAKONI CHOWDARI (29)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","RAKSHYALA B (30)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","DUPATAMAHAGANVA (31)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","MANIGEMPURA (32)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","LADHA (33)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","BACHEPALLI (34)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","BABALI (35)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","HEDAGAPURA (36)"],
        ["BIDAR (5)","AURAD (3)","SANTAPURA (3)","NITDORA K (37)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","RAKSHSHALA K (2)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","MASKALA (3)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","NAGURA B (4)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","BELAKONI CHOWDARI (5)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","RAKSHYALA B (6)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","DUPATAMAHAGANVA (10)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","MANIGEMPURA (11)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","LADHA (14)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","BACHEPALLI (15)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","BABALI (16)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","HEDAGAPURA (17)"],
        ["BIDAR (5)","AURAD (3)","TANAKUSANURA (4)","NITDORA K (18)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","BONTI (4)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","HANGARAGABI (5)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","HOKRANA (6)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","KHERDABI (7)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","KARAKYALA (8)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","DONGARAGAVA (9)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","BANDARA KUMATA (14)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","LINGI (15)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","BAVALAGANVA (20)"],
        ["BIDAR (5)","AURAD (3)","DABAKA (6)","SAVARAGANVA (21)"],
        // BIDAR TALUK
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","CHILLARGI (1)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","CHIMAKODA (2)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","KHAJAPURA (3)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","BASANTAPURA (4)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","KANGATI (5)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","FATTEPURACH (6)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","KAPALAPURAJH (7)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","JAMPADA (8)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","AMADALAPADA (9)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","ALAMASAPURA (10)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","RASULABADA (11)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","MALEGANVA (12)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","NANDAGANVA (13)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","MIRAJAPURAKH (14)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","GADAGI (15)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","HAMILAPURA (16)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","CHIKKAPETA (17)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","NAVADAGERI (18)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","GUMMA (19)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","SHAMARAJAPURA (20)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","SOLAPURA (21)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","SIPPALAGERA (22)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","MAMANAKERI (23)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","KABEERAVADA (24)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","TAJALAPURA (25)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","ODAVADA (26)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","ALIYABADA (27)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","IMAMABADA (28)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","FATTEPURAJH (29)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","SHERIHUSENI (30)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","BAVAPURA (31)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","DATTANAKERI (32)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","ALANKERI (33)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","KHILLARKA (34)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","SULTANAPURAKH (35)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARA UTTARA (1)","BAGE KARANJA (36)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","ASHTURA (1)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","MIRAJAPURAT (2)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","MALKAPURA (3)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","SULTANAPURAJH (4)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","HALADAKERI KH (5)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","MAILURA (6)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","KUMBARAVADA (7)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","AMALAPURA (8)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","GORANALLI BH (9)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","CHITTA (10)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","GUNNALLI (11)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","SHAHAPURA (12)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","HALADAKERI JH (13)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","BRAHMAPURA (14)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","BAGEHAMAMA (15)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","SHERICHAMPA (16)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","GULLARAHAVEELI (17)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","JANGALAKOYI (18)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","BAGEIBRAHIMA (19)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","MAILATEGAIRABI (20)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","BAGE-KHARIDI (21)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","MAILATEGORANALLI (22)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","BAGE-SHAHI (23)"],
        ["BIDAR (5)","BIDAR (4)","BEEDARADAKSHINA (2)","BAGEGORANALLI (24)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","MANNALLI (1)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","HOKRANABH (2)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","HOKRANAKH (3)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","NAGORA (4)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","GHODEPALLI (5)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","BUDHERA (6)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","T MIRJAPURA (7)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","CHINTALAGERA (8)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","DHARMAPURA (9)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","GOWSAPURA (10)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","YAKATAPURA (11)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","SATOLI (12)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","BARURA (13)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","RAJAGIRA (14)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","SINDHOLA (15)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","REKULAGI (16)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","TADAPALLI (17)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","BAMBALAGI (18)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","MARAKUNDA (19)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","BHANGURA (20)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","MUGADALA (21)"],
        ["BIDAR (5)","BIDAR (4)","MANNALLI (3)","BAPURA (22)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","BAGADALA (1)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","OWRADES (2)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","SIRASI A (3)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","HONNADDI (4)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","SHEKHAPURA (5)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","PATARAPALLI (6)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","BHAIRANALLI CH (7)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","KHASHEMPURACH (8)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","BAVAGI (9)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","NELAVADA (10)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","KHASHEMPURAP (11)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","MALIKA MEERJAPURA (12)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","MANDAKANALLI (13)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","KANGANAKOTA (14)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","KADAVADA (15)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","SHEMASHANAGARA (16)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","CHATNALLI (17)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","BARIDABADA (18)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","SANGOLAGI (19)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","RANJOLA  KHENI (20)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","NIDAVANCHA (21)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","SIRAKATANALLI (22)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","HOCHAKANALLI (23)"],
        ["BIDAR (5)","BIDAR (4)","BAGADALA (4)","HAJJARAGI (24)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","KAMATANA (1)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","SIKINDRAPURA (2)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","KHADARANAGARA (3)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","BAKACHOWDI (4)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","NIJAMAPURA (5)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","BELLURA (6)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","JAMISTANAPURA (7)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","YADLAPURA (8)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","KUTTABADA (9)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","AYASAPURA (10)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","NOWBADA (11)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","CHONDI (12)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","CHOWLI (13)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","KOLARAKE (14)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","CHIDRI (15)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","KOLARABI (16)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","ANADURA (17)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","KAPALAPURAA (18)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","ATIVALA (19)"],
        ["BIDAR (5)","BIDAR (4)","KAMATANA (5)","HONNIKERI (20)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","JANAVADA (1)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","KHANAJAMALAPURA (2)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","ALIYAMBARA (3)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","MAMADAPURA (4)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","YARANALLIPI (5)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","YARANALLIDI (6)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","BAMPALLI (7)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","ISLAMPURA (8)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","MARAKHALA (9)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","BENAKANALLI (10)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","DADDAPURA (11)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","SHREEMANDALA (12)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","NEMATABADA (13)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","NAVALASAPURA (14)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","CHAMBOLA (15)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","HIPPALAGANVA (16)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","KANNALLI (17)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","ALLAPURA (18)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","RAJANALA (19)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","SIDDAPURA (20)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","SANGAVI (21)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","VILASAPURA (22)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","BHAIRANALLIDI (23)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","GUMATAPURA (24)"],
        ["BIDAR (5)","BIDAR (4)","JANAVADA (6)","IMAMAPURA (25)"],
        // HUMNABAD TALUK
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","HUMANABADA (1)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","DHUMMANASURA (2)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","HUDAGI (3)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","GADAVANTI (4)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","MANIKANAGARA (5)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","MOLAKERA (6)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","HANAKUNI (7)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","BORAMPALLI (8)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","KAPPARAGANVA (9)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","NANDAGANVA (10)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","HALLIKHEDAKE (11)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","MUSTAPURA (12)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","CHITAKOTA (13)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","KALLURA (14)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","KATALLI (15)"],
        ["BIDAR (5)","HUMNABAD (5)","HUMANABADA (1)","SINDHANAKERA (16)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","HALLILIKHEDABI (1)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","SINDABANDAGI (2)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","ATIVALA (3)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","NAMADAPURA (4)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","BENACHINCHOLI (5)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","KABEERABADA (6)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","MADARAGANVA (7)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","MALKAPURA VADI (8)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","DAKULAGI (9)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","AMIRABADA (10)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","ALLURA (11)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","BOTAGI (12)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","MARAKHALA (13)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","HILALAPURA (14)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","SHAKKARAGANJA (15)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","SITALAGERA (16)"],
        ["BIDAR (5)","HUMNABAD (5)","HALLILIKHEDABI (2)","NIMBURA (17)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","DUBALAGUNDI (1)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","GHODAVADI (2)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","GHATABORALA (3)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","HUNASAGERA (4)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","KANAKATTA (5)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","HUNASANALA (6)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","SEDOLA (7)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","JALASINGI (8)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","KUMARA CHINCHOLI (9)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","HANDIKERA (10)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","VARAVATTIKE (11)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","OTAGI (12)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","SULTANABADA (13)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","CHEENAKERA (14)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","MUGANURA (15)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","CHANDANAHALLI (16)"],
        ["BIDAR (5)","HUMNABAD (5)","DUBALAGUNDI (3)","SONAKERA (17)"],
        // CHITAGUPPA TALUK
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","CHITAGUPPA (1)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","MUSTARI (2)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","GURUDALA (3)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","KODAMBALA (4)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","BELAKERA (5)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","VALAKHINDI (6)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","RAMAPURA (7)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","TALAMADAGI (8)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","HIPPARGA (9)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","MUDANALA (10)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","ITAGA (11)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","SHAMATABADA (12)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","KANDAGOLA (13)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","VADDANAKERA (14)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","MADAGOLA (15)"],
        ["BIDAR (5)","CHITAGUPPA (6)","CHITAGUPPA (4)","BHASKARA NAGARA (16)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","NIRNA (1)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","MADARAGI (2)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","BHADRAPURA (3)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","BASHEERAPURA (4)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","DEVAGEERI (5)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","MUTTANGI (6)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","ALIPURA (7)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","UDABALA (8)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","BANNALLI (9)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","MANGALAGI (10)"],
        ["BIDAR (5)","CHITAGUPPA (6)","NIRNA (5)","NAGANAKERA (11)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","BEMALAKHEDA (1)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","UDAMANALLI (2)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","POLAKAPALLI (3)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","KARAPAPALLI (4)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","KARAKANALLI (5)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","SAIDAPURA (6)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","MEENAKERA (7)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","BORALA (8)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","MANNAEKHELLI (9)"],
        ["BIDAR (5)","CHITAGUPPA (6)","BEMALAKHEDA (6)","CHANGALERA (10)"],
        // KAMALANAGARA TALUK
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","TANAKUSANURA (1)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","BALATA B (7)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","BALATA K (8)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","CHANDORI (9)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","BEDAKUNDA (12)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","HIPPALAGANVA (13)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","NIDODA (19)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","HALAHALLI (20)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","KOREKALA (21)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","BEMBRA (22)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","LINGADALLIYU (23)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","MUDHOLABI (24)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","MUDHOLA K (25)"],
        ["BIDAR (5)","KAMALANAGARA (7)","TANAKUSANURA (4)","SANGAMA (26)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","KAMALANAGARA (1)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","TORANA (2)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","DONAGANVA M (3)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","BAVANIBIJALAGANVA (4)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","SAVALI (5)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","BELAKUNIBHO (6)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","SONALA (7)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","HORANDI (8)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","BALURA K (9)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","MADANURA (10)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","MURGAKE (11)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","CHANDESHVARA (12)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","KHEDA (13)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","KHATAGANVA (14)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","RANDYALA (15)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","KALAGAPURA (16)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","HULASURAKHEDA (17)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","DIGGI (18)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","KOTAGYALA (19)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","RAMAPURA (20)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","TAPASHYALA (21)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","BASANALA (22)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","KORIYALA (23)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","HAKYALA (24)"],
        ["BIDAR (5)","KAMALANAGARA (7)","KAMALANAGARA (5)","HOLA SAMUDRA (25)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","DABAKA (1)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","VAGANAGERA (2)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","AKANAPURA (3)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","MALEGANVA (10)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","MUTAKHEDA (11)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","CHONDIMUKHEDA (12)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","NANDIBIJALAGANVA (13)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","CHIMEGANVA (16)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","HANDIKERA (17)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","CHIKALI U (18)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","GANESHAPURAYU (19)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","SANGANALA (22)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","DAREGANVA (23)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","MURKI (24)"],
        ["BIDAR (5)","KAMALANAGARA (7)","DABAKA (6)","GANGANABEEDA (25)"],
        // HULASURA TALUK
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","HULASURA (1)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","SOLADABAKA (2)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","TOGALURA (3)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","KADARABADA (4)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","DEVANALA (5)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","HALAHALLI (6)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","MACHANALA (7)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","GADIGOWNDAGANVA (8)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","MIRKALA (9)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","GUTTI (10)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","BELURA (11)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","GORTABH (12)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","MACHALAMBA (13)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","GADIRAYAPALLI (14)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","KADEPURA (15)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","LIMBAPURA (16)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","MUSTAPURA (17)"],
        ["BIDAR (5)","HULASURA (8)","HULASURA (6)","KOTAMALA (18)"]
    ];

    const talukMap = {
        "BASAVAKALYAN (1)":"BASAVAKALYAN (3)",
        "BHALKI (2)":"BHALKI (2)",
        "AURAD (3)":"AURAD (4)",
        "BIDAR (4)":"BIDAR (1)",
        "HUMNABAD (5)":"HUMNABAD (6)",
        "CHITAGUPPA (6)":"CHITAGUPPA (5)",
        "KAMALANAGARA (7)":"KAMALANAGARA (4)",
        "HULASURA (8)":"HULASURA (3)"
    };

    const rows = rawRowsOriginal.map(r => [r[0], talukMap[r[1]] || r[1], r[2], r[3]]);
    console.log('✅ Dataset loaded. Total rows:', rows.length);

    // ================================================================
    //  TABLE DEFINITIONS
    // ================================================================
    const tableDefs = [
        { value: 'hst_appl_mutation', group: 'application', label: 'hst_appl_mutation' },
        { value: 'appl_mutation', group: 'application', label: 'appl_mutation' },
        { value: 'appl_mutation_owners', group: 'application', label: 'appl_mutation_owners' },
        { value: 'tr_mutation', group: 'transaction', label: 'tr_mutation' },
        { value: 'tr_mutation_surveynos', group: 'transaction', label: 'tr_mutation_surveynos' },
        { value: 'tr_moldown', group: 'transaction', label: 'tr_moldown' },
        { value: 'tr_mnewown', group: 'transaction', label: 'tr_mnewown' },
        { value: 'tr_podi_rtc', group: 'transaction', label: 'tr_podi_rtc' },
        { value: 'tr_podi_owner', group: 'transaction', label: 'tr_podi_owner' },
        { value: 'tr_crop', group: 'transaction', label: 'tr_crop' },
        { value: 'tr_mcultiv', group: 'transaction', label: 'tr_mcultiv' },
        { value: 'hst_crop', group: 'history', label: 'hst_crop' },
        { value: 'hst_cultiv', group: 'history', label: 'hst_cultiv' },
        { value: 'hst_head', group: 'history', label: 'hst_head' },
        { value: 'hst_owner', group: 'history', label: 'hst_owner' },
        { value: 'rtc_head', group: 'rtc', label: 'rtc_head' },
        { value: 'rtc_owner', group: 'rtc', label: 'rtc_owner' },
        { value: 'rtc_cultiv', group: 'rtc', label: 'rtc_cultiv' },
        { value: 'rtc_crop', group: 'rtc', label: 'rtc_crop' },
        { value: 'GSC_SERVICES_TRAN', group: 'gsc', label: 'GSC_SERVICES_TRAN' },
        { value: 'GSC_UPDATE_TRAN', group: 'gsc', label: 'GSC_UPDATE_TRAN' }
    ];

    // ================================================================
    //  JAVASCRIPT LOGIC
    // ================================================================
    let currentChecked = new Set();

    function renderCheckboxes(groupFilter = 'all') {
        const container = document.getElementById('tableSelectContainer');
        const filtered = groupFilter === 'all' ? tableDefs : tableDefs.filter(t => t.group === groupFilter);
        let html = `<label class="all-tables"><input type="checkbox" id="selectAllTables"> ✅ All Tables</label>`;
        filtered.forEach(t => {
            const checked = currentChecked.has(t.value) ? 'checked' : '';
            html += `<label><input type="checkbox" value="${t.value}" ${checked}> ${t.label}</label>`;
        });
        container.innerHTML = html;

        const selectAll = document.getElementById('selectAllTables');
        if (selectAll) {
            selectAll.addEventListener('change', function() {
                const cbs = container.querySelectorAll('input[type="checkbox"]:not(#selectAllTables)');
                cbs.forEach(cb => cb.checked = this.checked);
                currentChecked.clear();
                cbs.forEach(cb => { if (cb.checked) currentChecked.add(cb.value); });
            });
        }
        container.querySelectorAll('input[type="checkbox"]:not(#selectAllTables)').forEach(cb => {
            cb.addEventListener('change', function() {
                if (this.checked) currentChecked.add(this.value);
                else currentChecked.delete(this.value);
                const allCbs = container.querySelectorAll('input[type="checkbox"]:not(#selectAllTables)');
                const allChecked = Array.from(allCbs).every(c => c.checked);
                const sa = document.getElementById('selectAllTables');
                if (sa) sa.checked = allChecked;
            });
        });
    }

    document.querySelectorAll('input[name="tableGroup"]').forEach(radio => {
        radio.addEventListener('change', function() {
            renderCheckboxes(this.value);
        });
    });

    renderCheckboxes('all');

    // ---- helpers ----
    function extractCode(str) {
        if (!str) return null;
        const m = str.match(/\((\d+)\)/);
        return m ? parseInt(m[1], 10) : null;
    }
    function escapeHtml(s) {
        if (!s) return '';
        return s.replace(/[&<>]/g, m => ({'&':'&amp;','<':'&lt;','>':'&gt;'})[m] || m);
    }
    function getUniqueValues(col, data) {
        const s = new Set();
        data.forEach(r => s.add(r[col]));
        return Array.from(s).sort();
    }

    // ---- DOM refs ----
    const districtSel = document.getElementById('districtSelect');
    const talukSel = document.getElementById('talukSelect');
    const hobliSel = document.getElementById('hobliSelect');
    const villageSel = document.getElementById('villageSelect');
    const searchInput = document.getElementById('villageSearch');
    const tbody = document.getElementById('tableBody');
    const totalSpan = document.getElementById('totalCount');
    const showSpan = document.getElementById('showCount');
    const codeStats = document.getElementById('codeStats');
    const queryGrid = document.getElementById('queryGrid');
    const clearBtn = document.getElementById('clearAllBtn');
    const submitBtn = document.getElementById('submitBtn');
    const applNoInput = document.getElementById('applNoInput');
    const tranNoInput = document.getElementById('tranNoInput');
    const landCodeInput = document.getElementById('landCodeInput');
    const yearCodeInput = document.getElementById('yearCodeInput');

    // ---- Populate dropdowns ----
    function populateDistricts() {
        const vals = getUniqueValues(0, rows);
        districtSel.innerHTML = '<option value="">All Districts</option>' +
            vals.map(v => `<option value="${escapeHtml(v)}">${escapeHtml(v)}</option>`).join('');
    }

    function populateTaluks() {
        const dist = districtSel.value;
        const filtered = dist ? rows.filter(r => r[0] === dist) : rows;
        const vals = getUniqueValues(1, filtered);
        talukSel.innerHTML = '<option value="">All Taluks</option>' +
            vals.map(v => `<option value="${escapeHtml(v)}">${escapeHtml(v)}</option>`).join('');
        populateHoblis();
    }

    function populateHoblis() {
        const dist = districtSel.value;
        const taluk = talukSel.value;
        let filtered = rows;
        if (dist) filtered = filtered.filter(r => r[0] === dist);
        if (taluk) filtered = filtered.filter(r => r[1] === taluk);
        const vals = getUniqueValues(2, filtered);
        hobliSel.innerHTML = '<option value="">All Hoblis</option>' +
            vals.map(v => `<option value="${escapeHtml(v)}">${escapeHtml(v)}</option>`).join('');
        populateVillages();
    }

    function populateVillages() {
        const dist = districtSel.value;
        const taluk = talukSel.value;
        const hobli = hobliSel.value;
        let filtered = rows;
        if (dist) filtered = filtered.filter(r => r[0] === dist);
        if (taluk) filtered = filtered.filter(r => r[1] === taluk);
        if (hobli) filtered = filtered.filter(r => r[2] === hobli);
        const vals = getUniqueValues(3, filtered);
        villageSel.innerHTML = '<option value="">All Villages</option>' +
            vals.map(v => `<option value="${escapeHtml(v)}">${escapeHtml(v)}</option>`).join('');
        applyFilters();
    }

    // ---- Filter table ----
    function applyFilters() {
        const dist = districtSel.value;
        const taluk = talukSel.value;
        const hobli = hobliSel.value;
        const village = villageSel.value;
        const term = searchInput.value.trim().toLowerCase();

        const filtered = rows.filter(r => {
            if (dist && r[0] !== dist) return false;
            if (taluk && r[1] !== taluk) return false;
            if (hobli && r[2] !== hobli) return false;
            if (village && r[3] !== village) return false;
            if (term && !r[3].toLowerCase().includes(term)) return false;
            return true;
        });

        renderTable(filtered);
        totalSpan.textContent = rows.length;
        showSpan.textContent = filtered.length;
        updateCodeStats();
    }

    function renderTable(data) {
        if (!data.length) {
            tbody.innerHTML = `<tr><td colspan="4" class="no-results">No records found</td></tr>`;
            return;
        }
        let html = '';
        data.forEach(r => {
            const vname = escapeHtml(r[3]);
            html += `<tr>
                <td>${escapeHtml(r[0])}</td>
                <td>${escapeHtml(r[1])}</td>
                <td>${escapeHtml(r[2])}</td>
                <td><span class="village-link" data-dist="${escapeHtml(r[0])}" data-taluk="${escapeHtml(r[1])}" data-hobli="${escapeHtml(r[2])}" data-village="${escapeHtml(r[3])}">${vname}</span></td>
            </tr>`;
        });
        tbody.innerHTML = html;

        // --- FIXED: Click handler for village links ---
        document.querySelectorAll('.village-link').forEach(el => {
            el.addEventListener('click', function(e) {
                const dist = this.dataset.dist;
                const taluk = this.dataset.taluk;
                const hobli = this.dataset.hobli;
                const village = this.dataset.village;

                // Set district and rebuild taluks
                districtSel.value = dist;
                const filteredTaluks = rows.filter(r => r[0] === dist);
                const taluks = getUniqueValues(1, filteredTaluks);
                talukSel.innerHTML = '<option value="">All Taluks</option>' +
                    taluks.map(t => `<option value="${escapeHtml(t)}">${escapeHtml(t)}</option>`).join('');
                talukSel.value = taluk;

                // Rebuild hoblis based on district + taluk
                const filteredHoblis = rows.filter(r => r[0] === dist && r[1] === taluk);
                const hoblis = getUniqueValues(2, filteredHoblis);
                hobliSel.innerHTML = '<option value="">All Hoblis</option>' +
                    hoblis.map(h => `<option value="${escapeHtml(h)}">${escapeHtml(h)}</option>`).join('');
                hobliSel.value = hobli;

                // Rebuild villages based on district + taluk + hobli
                const filteredVillages = rows.filter(r => r[0] === dist && r[1] === taluk && r[2] === hobli);
                const villages = getUniqueValues(3, filteredVillages);
                villageSel.innerHTML = '<option value="">All Villages</option>' +
                    villages.map(v => `<option value="${escapeHtml(v)}">${escapeHtml(v)}</option>`).join('');
                villageSel.value = village;

                // Update search input and table
                searchInput.value = village;
                applyFilters();
                generateQueries();
            });
        });
    }

    // ---- Code stats ----
    function updateCodeStats() {
        const d = districtSel.value;
        const t = talukSel.value;
        const h = hobliSel.value;
        const v = villageSel.value;
        const dc = extractCode(d);
        const tc = extractCode(t);
        const hc = extractCode(h);
        const vc = extractCode(v);
        let parts = [];
        if (dc !== null) parts.push(`dist_code=${dc}`);
        if (tc !== null) parts.push(`taluk_code=${tc}`);
        if (hc !== null) parts.push(`hobli_code=${hc}`);
        if (vc !== null) parts.push(`village_code=${vc}`);
        codeStats.textContent = parts.length ? `📌 ${parts.join(' AND ')}` : '📌 Codes: —';
        return { dc, tc, hc, vc };
    }

    // ---- Generate Queries ----
    function generateQueries() {
        const d = districtSel.value;
        const t = talukSel.value;
        const h = hobliSel.value;
        const v = villageSel.value;
        const appl = applNoInput.value.trim();
        const tran = tranNoInput.value.trim();
        const land = landCodeInput.value.trim();
        const year = yearCodeInput.value.trim();

        const dc = extractCode(d);
        const tc = extractCode(t);
        const hc = extractCode(h);
        const vc = extractCode(v);

        const checked = Array.from(currentChecked);
        if (!checked.length) {
            queryGrid.innerHTML = '⚠️ Please select at least one table.';
            return;
        }

        const conds = [];
        if (dc !== null) conds.push(`dist_code = ${dc}`);
        if (tc !== null) conds.push(`taluk_code = ${tc}`);
        if (hc !== null) conds.push(`hobli_code = ${hc}`);
        if (vc !== null) conds.push(`village_code = ${vc}`);
        if (appl) conds.push(`appl_no = '${appl}'`);
        if (tran) conds.push(`tran_no = '${tran}'`);
        if (land) conds.push(`land_code = '${land}'`);
        if (year) conds.push(`year_code = '${year}'`);

        if (conds.length === 0) {
            queryGrid.innerHTML = '⚠️ No filters selected. Please select at least one filter (location, app no, etc.)';
            return;
        }

        const where = conds.join(' AND ');
        let html = '';
        checked.forEach(tbl => {
            html += `SELECT * FROM ${tbl} WHERE ${where};\\n`;
        });
        queryGrid.innerHTML = html;

        // Auto-fill report form
        const issueDesc = `Context: ${d || 'All'} → ${t || 'All'} → ${h || 'All'} → ${v || 'All'}\\nCodes: ${conds.join(' AND ')}`;
        document.getElementById('form_issue_desc').value = issueDesc;
        document.getElementById('form_select_block').value = html;
        if (!document.getElementById('reportAppNo').value) document.getElementById('reportAppNo').value = appl;
        if (!document.getElementById('reportTranNo').value) document.getElementById('reportTranNo').value = tran;
        if (!document.getElementById('reportYear').value) document.getElementById('reportYear').value = year;
        if (!document.getElementById('reportLandCode').value) document.getElementById('reportLandCode').value = land;
    }

    // ---- Reset ----
    function resetAll() {
        districtSel.value = '';
        talukSel.innerHTML = '<option value="">All Taluks</option>';
        hobliSel.innerHTML = '<option value="">All Hoblis</option>';
        villageSel.innerHTML = '<option value="">All Villages</option>';
        searchInput.value = '';
        applNoInput.value = '';
        tranNoInput.value = '';
        landCodeInput.value = '';
        yearCodeInput.value = '';
        document.getElementById('reportAppNo').value = '';
        document.getElementById('reportTranNo').value = '';
        document.getElementById('reportYear').value = '';
        document.getElementById('reportLandCode').value = '';
        document.getElementById('reportSeason').value = '';
        document.getElementById('form_issue_desc').value = '';
        document.getElementById('form_select_block').value = '';
        document.getElementById('form_update_block').value = '';
        currentChecked.clear();
        document.querySelectorAll('#tableSelectContainer input[type="checkbox"]').forEach(cb => cb.checked = false);
        const sa = document.getElementById('selectAllTables');
        if (sa) sa.checked = false;
        populateTaluks();
        queryGrid.innerHTML = 'Select filters and click Generate Queries';
        codeStats.textContent = '📌 Codes: —';
    }

    // ---- Event listeners ----
    districtSel.addEventListener('change', populateTaluks);
    talukSel.addEventListener('change', populateHoblis);
    hobliSel.addEventListener('change', populateVillages);
    villageSel.addEventListener('change', applyFilters);
    searchInput.addEventListener('input', applyFilters);
    clearBtn.addEventListener('click', resetAll);
    submitBtn.addEventListener('click', function() {
        generateQueries();
        applyFilters();
    });

    // ---- Init ----
    populateDistricts();
    populateTaluks();
    applyFilters();
    console.log('✅ Initialization complete.');
</script>
</body>
</html>
"""

@app.route('/')
def home():
    today_str = datetime.now().strftime('%d-%m-%Y')
    return render_template_string(HTML_TEMPLATE, current_date=today_str)

@app.route('/generate', methods=['POST'])
def generate():
    form_data = {
        'report_date': request.form.get('report_date', datetime.now().strftime('%d-%m-%Y')),
        'cls_number': request.form.get('cls_number', ' '),
        'app_no': request.form.get('app_no', ''),
        'tran_no': request.form.get('tran_no', ''),
        'year': request.form.get('year', ''),
        'land_code': request.form.get('land_code', ''),
        'season': request.form.get('season', ''),
        'issue_type': request.form.get('issue_type', 'Technical'),
        'issue_desc': request.form.get('issue_desc', ''),
        'findings_desc': request.form.get('findings_desc', ''),
        'root_cause': request.form.get('root_cause', ''),
        'sol_provided': request.form.get('sol_provided', ''),
        'select_block': request.form.get('select_block', ''),
        'update_block': request.form.get('update_block', '')
    }
    doc_stream = generate_docx(form_data)
    return send_file(
        doc_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='Issue_Analysis_and_Resolution_Report.docx'
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('FLASK_ENV') != 'production'
    print("--------------------------------------------------")
    print(f" 🚀 Server running on port {port}")
    print(" ✅ Village click fixed – dropdowns update correctly.")
    print("--------------------------------------------------")
    app.run(host='0.0.0.0', port=port, debug=debug_mode)